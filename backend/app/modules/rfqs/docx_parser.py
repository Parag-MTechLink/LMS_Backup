import docx
import re
from typing import Dict, Any, Optional
from datetime import datetime

def format_date(date_str: str) -> str:
    """Converts DD-MM-YYYY or other common formats to YYYY-MM-DD."""
    if not date_str:
        return ""
    
    # Try DD-MM-YYYY first (user's format)
    try:
        dt = datetime.strptime(date_str, "%d-%m-%Y")
        return dt.strftime("%Y-%m-%d")
    except ValueError:
        pass
        
    # Try DD/MM/YYYY
    try:
        dt = datetime.strptime(date_str, "%d/%m/%Y")
        return dt.strftime("%Y-%m-%d")
    except ValueError:
        pass
        
    # Fallback: if it's already YYYY-MM-DD, just return it
    if re.match(r"\d{4}-\d{2}-\d{2}", date_str):
        return date_str
        
    return date_str

def parse_rfq_docx(file_stream) -> Dict[str, Any]:
    """Parses the RFQ/JRF docx file and extracts key fields."""
    doc = docx.Document(file_stream)
    
    # We mainly care about Table 0
    if not doc.tables:
        return {}
        
    table = doc.tables[0]
    
    def get_cell_text(row, col):
        try:
            return table.rows[row].cells[col].text.strip()
        except (IndexError, AttributeError):
            return ""

    # Extract fields based on identified mapping in Job Request Form - 1.docx
    
    # R4 C0: Ref No: 123ABC-XYZ321\nDate: 28-03-2026
    ref_date_text = get_cell_text(4, 0)
    ref_no = ""
    received_date = ""
    if "Ref No:" in ref_date_text:
        parts = ref_date_text.split("\n")
        for p in parts:
            if "Ref No:" in p:
                ref_no = p.replace("Ref No:", "").strip()
            if "Date:" in p:
                raw_date = p.replace("Date:", "").strip()
                received_date = format_date(raw_date)

    # R7 C1: Company Name
    customer_name = get_cell_text(7, 1)
    
    # R8 C1: Contact Person, R8 C9: Designation
    contact_person = get_cell_text(8, 1)
    designation = get_cell_text(8, 9)
    
    # R9 C1: Address, R9 C9: Email
    address = get_cell_text(9, 1)
    email = get_cell_text(9, 9)
    
    # R10 C1: Telephone
    telephone_text = get_cell_text(10, 1)
    telephone = telephone_text.replace("Telephone:", "").strip()
    
    # R11 C1: Mobile
    mobile_text = get_cell_text(11, 1)
    mobile = mobile_text.replace("Mobile:", "").strip()
    
    # R14 C9: Project Code, R14 C15: Task
    project_code = get_cell_text(14, 9)
    # Handle scientific notation if it comes through as such (e.g. 5.67891E+11)
    try:
        if "E" in project_code.upper():
            project_code = str(int(float(project_code)))
    except:
        pass
    
    # R23 C2: Product Name
    product = get_cell_text(23, 2)
    # R23 C10: Supply Voltage
    supply_voltage = get_cell_text(23, 10)
    
    # R24 C2: Quantity
    quantity = get_cell_text(24, 2)
    # R24 C10: Operating Frequency
    operating_freq = get_cell_text(24, 10)
    
    # R25 C2: Manufacturer
    manufacturer = get_cell_text(25, 2)
    # R25 C10: Current
    current = get_cell_text(25, 10)
    
    # R26 C2: Model No
    model_no = get_cell_text(26, 2)
    # R26 C10: Weight
    weight = get_cell_text(26, 10)
    
    # R27 C2: Serial No
    serial_no = get_cell_text(27, 2)
    # R27 C10: Dimensions
    dimensions = get_cell_text(27, 10)

    # R28 C2: Power Ports
    power_ports = get_cell_text(28, 2)
    # R28 C10: Signal Lines
    signal_lines = get_cell_text(28, 10)

    return {
        "customer": {
            "companyName": customer_name,
            "contactPerson": contact_person,
            "designation": designation,
            "address": address,
            "email": email,
            "telephone": telephone,
            "mobile": mobile
        },
        "rfq": {
            "product": product,
            "receivedDate": received_date,
            "refNo": ref_no,
            "projectCode": project_code,
            "quantity": quantity,
            "manufacturer": manufacturer,
            "modelNo": model_no,
            "serialNo": serial_no,
            "raw_details": {
               "Designation": designation,
               "Project Code": project_code,
               "Quantity": quantity,
               "Manufacturer": manufacturer,
               "Model No": model_no,
               "Serial No": serial_no,
               "Supply Voltage": supply_voltage,
               "Operating Frequency": operating_freq,
               "Current": current,
               "Weight": weight,
               "Dimensions": dimensions,
               "Power Ports": power_ports,
               "Signal Lines": signal_lines
            }
        }
    }
